import os, re, pyperclip
import win32com.client
from datetime import datetime
from docx import Document

# Boas vindas
print("Bem vindo ao AUTOCOTA!\n")

# Configurações do usuário
def criar_diretorio(diretorio):
    if not os.path.exists(diretorio):
        os.makedirs(diretorio)

criar_diretorio('docx_gerados')
criar_diretorio('pdfs_gerados')

docx_folder = r'C:\Users\jepim\Desktop\autocota\docx_gerados'
pdf_folder = r'C:\Users\jepim\Desktop\autocota\pdfs_gerados' 

class DocumentProcessor:
    def __init__(self):
        self.data = datetime.now().strftime('%Y-%m-%d')

    def process_documents(self):
        with open('cotas.txt', "r", encoding='utf-8') as file:
            texto = file.read()

        resultados = texto.split('-------------------------\n')

        for resultado in resultados:
            parametros = resultado.split('\n')
            numero_do_processo = parametros[0]
            texto_da_cota = []
            for parametro in parametros:
                if parametro != parametros[0] and parametro != '':
                    texto_da_cota.append(parametro)

            if numero_do_processo != '' and texto_da_cota != []:
                doc = Document('documento_base.docx')
                self.fill_process_number(doc, numero_do_processo)
                self.insert_manifestation(doc, texto_da_cota)
                self.delete_mark(doc)
                self.save_document(doc, numero_do_processo)

    def fill_process_number(self, doc, process_number):
        for paragrafo in doc.paragraphs:
            paragrafo.text = paragrafo.text.replace('XXXX', process_number)

    def insert_manifestation(self, doc, texto_da_cota):
        for texto in texto_da_cota:
            for i, paragrafo in enumerate(doc.paragraphs):
                if 'YYYY' in paragrafo.text:
                    prior_paragraph = doc.paragraphs[i]
                    prior_paragraph.insert_paragraph_before(texto)

    def delete_mark(self, doc):
        for paragrafo in doc.paragraphs:
            paragrafo.text = paragrafo.text.replace('YYYY', '')

    def save_document(self, doc, process_number):
        folder_path = fr'.\docx_gerados'
        os.makedirs(folder_path, exist_ok=True)
        file_name = f'{process_number} - {self.data}.docx'
        file_path = os.path.join(folder_path, file_name)
        doc.save(file_path)

class PDFConverter:
    def __init__(self):
        self.word = win32com.client.Dispatch('Word.Application')
       
    def convert_to_pdf(self):
        for filename in os.listdir(docx_folder):
            if filename.endswith('.docx'):
                docx_path = os.path.join(docx_folder, filename)
                doc = self.word.Documents.Open(docx_path)
                pdf_name = os.path.splitext(filename)[0] + '.pdf'
                pdf_path = os.path.join(pdf_folder, pdf_name)
                doc.SaveAs(pdf_path, FileFormat=17)
                doc.Close()

def gather_intimations():
    resposta = input("As intimações a receber estão na área de transferência? (S/N): ")

    if resposta.lower() == "s":
        clipboard_text = pyperclip.paste()
    else:
        os._exit(0)
    
    user_input = clipboard_text
    text = user_input.replace('\n', '')
    padrao_processo = re.compile(r'\d{7}-\d{2}.\d{4}.8.26.\d{4}')
    processos_encontrados = re.findall(padrao_processo, text)

    with open('cotas.txt', "w") as file:
        for processo in processos_encontrados:
            file.write(processo + '\n\n')
            file.write('-------------------------\n')

    print("Arquivo 'cotas.txt' gerado com sucesso!")

def main():
    print("Escolha uma opção:")
    print("1. Gerar arquivo 'cotas.txt'")
    print("2. Gerar arquivos Word e PDF")

    opcao = input("Opção: ")

    if opcao == '1':
        gather_intimations()
    elif opcao == '2':
        document_processor = DocumentProcessor()
        document_processor.process_documents()

        pdf_converter = PDFConverter()
        pdf_converter.convert_to_pdf()
    else:
        print("Opção inválida.")

if __name__ == '__main__':
    main()